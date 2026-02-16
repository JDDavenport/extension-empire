#!/usr/bin/env python3
"""Create all 5 Etsy digital product files."""
import os
BASE = os.path.expanduser("~/clawd/projects/extension-empire/etsy-empire")

# ============================================================
# 1. FREELANCER SOW BUNDLE
# ============================================================
def create_sow_templates():
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    
    outdir = os.path.join(BASE, "freelancer-sow-bundle", "files")
    os.makedirs(outdir, exist_ok=True)
    
    industries = {
        "web-development": {
            "title": "Web Development",
            "scope_items": [
                "Design and develop a responsive website with up to [X] pages",
                "Implement content management system (CMS) integration",
                "Set up hosting environment and SSL certificate",
                "Cross-browser and mobile device testing",
                "Search engine optimization (SEO) basics implementation",
                "Integration of analytics tracking (Google Analytics)",
                "Contact form and email notification setup",
            ],
            "deliverables": [
                ("Website Design Mockups", "Figma/Adobe XD files with desktop and mobile layouts"),
                ("Front-End Development", "HTML5, CSS3, JavaScript responsive code"),
                ("Back-End Development", "CMS setup, database configuration, server-side logic"),
                ("Testing & QA Report", "Cross-browser testing results, performance benchmarks"),
                ("Documentation", "Technical docs, login credentials, maintenance guide"),
                ("Training Session", "1-hour walkthrough of CMS and website management"),
            ],
            "milestones": [
                ("Discovery & Planning", "Week 1-2", "15%"),
                ("Design Mockups & Approval", "Week 3-4", "20%"),
                ("Front-End Development", "Week 5-7", "25%"),
                ("Back-End & CMS Integration", "Week 8-10", "25%"),
                ("Testing, QA & Launch", "Week 11-12", "15%"),
            ],
            "rate": "$75-150/hour or $5,000-25,000 fixed",
        },
        "graphic-design": {
            "title": "Graphic Design",
            "scope_items": [
                "Create brand identity package including logo, color palette, and typography",
                "Design [X] marketing collateral pieces (business cards, letterhead, etc.)",
                "Develop social media graphic templates",
                "Create print-ready files in required formats (CMYK, bleed, trim)",
                "Provide all source files in editable formats",
                "Up to [X] rounds of revisions per deliverable",
            ],
            "deliverables": [
                ("Primary Logo", "Full color, black/white, and reversed versions in AI, EPS, PNG, SVG"),
                ("Brand Style Guide", "Colors (Pantone, CMYK, RGB, HEX), typography, usage rules"),
                ("Business Card Design", "Front and back, print-ready PDF with bleed"),
                ("Letterhead & Envelope", "Branded templates in editable and print-ready formats"),
                ("Social Media Templates", "Profile images, cover photos, post templates for 3 platforms"),
                ("Marketing Materials", "Brochure/flyer designs as specified in project brief"),
            ],
            "milestones": [
                ("Creative Brief & Research", "Week 1", "10%"),
                ("Initial Concepts (3 directions)", "Week 2-3", "25%"),
                ("Refinement & Revisions", "Week 4-5", "30%"),
                ("Final Deliverables & File Prep", "Week 6", "25%"),
                ("Handoff & Style Guide", "Week 7", "10%"),
            ],
            "rate": "$50-125/hour or $2,500-15,000 fixed",
        },
        "marketing": {
            "title": "Marketing & Advertising",
            "scope_items": [
                "Develop comprehensive marketing strategy and campaign plan",
                "Create and manage paid advertising campaigns (Google Ads, Meta Ads)",
                "Produce content calendar and marketing copy",
                "Set up email marketing automation sequences",
                "Monthly performance reporting and optimization",
                "Competitor analysis and market positioning",
            ],
            "deliverables": [
                ("Marketing Strategy Document", "Target audience, positioning, channel strategy, KPIs"),
                ("Campaign Creative Assets", "Ad copy, images, landing page wireframes"),
                ("Content Calendar", "30/60/90-day content plan across all channels"),
                ("Email Sequences", "Welcome series, nurture sequence, re-engagement campaigns"),
                ("Monthly Reports", "Performance dashboards, insights, optimization recommendations"),
                ("Campaign Setup", "Fully configured ad accounts, tracking pixels, conversion goals"),
            ],
            "milestones": [
                ("Audit & Strategy Development", "Week 1-2", "20%"),
                ("Campaign Setup & Creative", "Week 3-4", "25%"),
                ("Launch & Initial Optimization", "Week 5-6", "20%"),
                ("Ongoing Management (Month 2)", "Week 7-10", "20%"),
                ("Reporting & Handoff", "Week 11-12", "15%"),
            ],
            "rate": "$2,000-10,000/month retainer",
        },
        "photography": {
            "title": "Photography",
            "scope_items": [
                "Pre-production planning including shot list and location scouting",
                "Professional photography session of [X] hours duration",
                "Post-production editing and color correction of selected images",
                "Delivery of high-resolution and web-optimized files",
                "Usage rights as specified in licensing section below",
                "Travel to location(s) within [X] mile radius",
            ],
            "deliverables": [
                ("Shot List & Creative Brief", "Approved list of required shots and creative direction"),
                ("Raw Selects Gallery", "Online gallery of unedited selects for client review"),
                ("Edited Final Images", "[X] fully retouched images in high-res TIFF and web-ready JPEG"),
                ("Behind-the-Scenes", "[X] BTS shots for social media use"),
                ("Print-Ready Files", "CMYK conversions for any images designated for print use"),
                ("Digital Delivery", "Cloud-based download link with organized folder structure"),
            ],
            "milestones": [
                ("Pre-Production & Planning", "Week 1", "25%"),
                ("Photo Shoot Day(s)", "Week 2", "25%"),
                ("Selects & Client Review", "Week 3", "0%"),
                ("Post-Production & Editing", "Week 4-5", "40%"),
                ("Final Delivery", "Week 6", "10%"),
            ],
            "rate": "$200-500/hour or $1,500-10,000 per project",
        },
        "consulting": {
            "title": "Business Consulting",
            "scope_items": [
                "Conduct comprehensive business assessment and needs analysis",
                "Develop strategic recommendations and implementation roadmap",
                "Facilitate [X] workshops/strategy sessions with stakeholders",
                "Provide industry benchmarking and best practices analysis",
                "Create standard operating procedures (SOPs) for key processes",
                "Deliver executive summary and board-ready presentation",
            ],
            "deliverables": [
                ("Discovery Report", "Current state assessment, SWOT analysis, gap identification"),
                ("Strategic Roadmap", "Prioritized recommendations with timeline and resource requirements"),
                ("Workshop Facilitation", "[X] half-day sessions with key stakeholders"),
                ("SOP Documentation", "Detailed procedures for [X] identified processes"),
                ("Implementation Guide", "Step-by-step action plan with owners and deadlines"),
                ("Executive Presentation", "Board-ready deck summarizing findings and recommendations"),
            ],
            "milestones": [
                ("Discovery & Assessment", "Week 1-3", "25%"),
                ("Analysis & Strategy Development", "Week 4-6", "25%"),
                ("Workshop Facilitation", "Week 7-8", "20%"),
                ("Documentation & Deliverables", "Week 9-11", "20%"),
                ("Final Presentation & Handoff", "Week 12", "10%"),
            ],
            "rate": "$150-350/hour or $10,000-50,000 per engagement",
        },
        "writing": {
            "title": "Content Writing & Copywriting",
            "scope_items": [
                "Research and write [X] pieces of content per month",
                "SEO keyword research and optimization for all content",
                "Develop brand voice guidelines and content style guide",
                "Create website copy for [X] pages",
                "Write email marketing copy for [X] campaigns",
                "Up to [X] rounds of revisions per piece",
            ],
            "deliverables": [
                ("Content Strategy", "Editorial calendar, keyword map, content pillars"),
                ("Blog Posts/Articles", "[X] posts of [X] words each, SEO-optimized"),
                ("Website Copy", "Homepage, About, Services, and [X] additional pages"),
                ("Email Copy", "[X] email sequences with subject lines and body copy"),
                ("Style Guide", "Brand voice document with tone, vocabulary, and examples"),
                ("Content Audit", "Analysis of existing content with improvement recommendations"),
            ],
            "milestones": [
                ("Research & Strategy", "Week 1-2", "20%"),
                ("First Draft Delivery", "Week 3-4", "30%"),
                ("Revisions & Refinement", "Week 5-6", "25%"),
                ("Final Copy & Style Guide", "Week 7-8", "15%"),
                ("Handoff & Training", "Week 8", "10%"),
            ],
            "rate": "$0.15-0.50/word or $2,000-8,000/month retainer",
        },
        "video-production": {
            "title": "Video Production",
            "scope_items": [
                "Pre-production: concept development, scriptwriting, storyboarding",
                "Production: filming with professional crew and equipment",
                "Post-production: editing, color grading, sound design, motion graphics",
                "Provide [X] rounds of revision cuts",
                "Deliver in multiple formats for web, social, and broadcast",
                "Music licensing and talent release coordination",
            ],
            "deliverables": [
                ("Creative Brief & Script", "Approved concept, script, and storyboard"),
                ("Production Day(s)", "[X] days of filming with crew, equipment, lighting"),
                ("Rough Cut", "Initial edit for client review and feedback"),
                ("Final Cut", "Polished edit with color grade, sound mix, graphics"),
                ("Alternate Versions", "Social media cuts (15s, 30s, 60s) for each platform"),
                ("Master Files", "ProRes/H.265 master, plus H.264 web versions, SRT captions"),
            ],
            "milestones": [
                ("Pre-Production & Planning", "Week 1-3", "25%"),
                ("Production/Filming", "Week 4-5", "30%"),
                ("Rough Cut & Review", "Week 6-7", "15%"),
                ("Final Edit & Color/Sound", "Week 8-9", "20%"),
                ("Delivery & Alternate Cuts", "Week 10", "10%"),
            ],
            "rate": "$2,000-20,000+ per finished minute",
        },
        "app-development": {
            "title": "App Development",
            "scope_items": [
                "Design and develop mobile application for [iOS/Android/Both]",
                "UI/UX design with interactive prototypes",
                "Backend API development and database architecture",
                "Third-party API integrations (payment, auth, analytics)",
                "App store submission and approval process management",
                "Post-launch bug fixes for [X] days",
            ],
            "deliverables": [
                ("Technical Architecture", "System design, tech stack, database schema, API spec"),
                ("UI/UX Designs", "Wireframes, high-fidelity mockups, interactive prototype"),
                ("MVP Application", "Fully functional app with core features"),
                ("Backend/API", "Scalable server infrastructure, API endpoints, documentation"),
                ("Testing Suite", "Unit tests, integration tests, QA test results"),
                ("App Store Package", "Submitted build, store listing, screenshots, metadata"),
            ],
            "milestones": [
                ("Discovery & Architecture", "Week 1-3", "15%"),
                ("UI/UX Design & Prototype", "Week 4-6", "15%"),
                ("Core Development Sprint 1", "Week 7-10", "25%"),
                ("Core Development Sprint 2", "Week 11-14", "25%"),
                ("QA, Testing & Launch", "Week 15-18", "20%"),
            ],
            "rate": "$100-200/hour or $25,000-150,000+ fixed",
        },
        "seo": {
            "title": "SEO (Search Engine Optimization)",
            "scope_items": [
                "Comprehensive technical SEO audit of existing website",
                "Keyword research and competitive analysis",
                "On-page optimization for [X] priority pages",
                "Link building and off-page SEO strategy execution",
                "Monthly performance tracking and reporting",
                "Local SEO optimization (if applicable)",
            ],
            "deliverables": [
                ("Technical SEO Audit", "Crawl report, site speed analysis, mobile-friendliness, schema markup"),
                ("Keyword Strategy", "Target keyword list with volume, difficulty, and intent mapping"),
                ("On-Page Optimization", "Title tags, meta descriptions, headers, internal linking for [X] pages"),
                ("Content Recommendations", "Content gap analysis with topic briefs for new content"),
                ("Link Building Report", "Backlink acquisition log, DA improvements, outreach results"),
                ("Monthly Analytics Report", "Rankings, traffic, conversions, competitive benchmarks"),
            ],
            "milestones": [
                ("Audit & Research", "Month 1", "25%"),
                ("On-Page Optimization", "Month 2", "25%"),
                ("Content & Link Building", "Month 3-4", "25%"),
                ("Ongoing Optimization", "Month 5-6", "15%"),
                ("Final Report & Handoff", "Month 6", "10%"),
            ],
            "rate": "$1,000-5,000/month retainer",
        },
        "social-media": {
            "title": "Social Media Management",
            "scope_items": [
                "Manage [X] social media platforms (e.g., Instagram, Facebook, LinkedIn, TikTok)",
                "Create and schedule [X] posts per week per platform",
                "Develop monthly content calendars with client approval",
                "Community management and engagement (comments, DMs, mentions)",
                "Paid social media advertising management",
                "Monthly analytics reporting and strategy adjustments",
            ],
            "deliverables": [
                ("Social Media Strategy", "Platform-specific strategy, audience personas, content pillars"),
                ("Content Calendar", "Monthly calendar with post copy, hashtags, and visuals"),
                ("Original Content", "[X] posts per week including graphics, captions, stories"),
                ("Community Management", "Daily monitoring, response within [X] hours, escalation protocol"),
                ("Ad Campaigns", "[X] paid campaigns per month with targeting and creative"),
                ("Monthly Report", "Follower growth, engagement rates, top posts, ROI analysis"),
            ],
            "milestones": [
                ("Audit & Strategy", "Week 1-2", "20%"),
                ("Content Creation & Calendar", "Week 3-4", "20%"),
                ("Month 1 Management", "Week 5-8", "20%"),
                ("Month 2 Management", "Week 9-12", "20%"),
                ("Month 3 & Handoff", "Week 13-16", "20%"),
            ],
            "rate": "$1,500-6,000/month retainer",
        },
    }
    
    for slug, info in industries.items():
        doc = Document()
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        
        # Title
        title = doc.add_heading('STATEMENT OF WORK', level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in title.runs:
            run.font.color.rgb = RGBColor(0x1a, 0x3c, 0x6e)
        
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run(f'{info["title"]} Services')
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0x2c, 0x5f, 0x8a)
        
        doc.add_paragraph()  # spacer
        
        # Info table
        table = doc.add_table(rows=6, cols=2)
        table.style = 'Light Shading Accent 1'
        cells = [
            ("Document ID:", "SOW-[YYYY]-[###]"),
            ("Client:", "[Client Name / Company]"),
            ("Contractor:", "[Your Name / Company]"),
            ("Prepared Date:", "[Date]"),
            ("Valid Until:", "[Date + 30 days]"),
            ("Version:", "1.0"),
        ]
        for i, (label, value) in enumerate(cells):
            table.rows[i].cells[0].text = label
            table.rows[i].cells[1].text = value
            for cell in table.rows[i].cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.size = Pt(10)
        
        doc.add_page_break()
        
        # Section 1: Project Overview
        doc.add_heading('1. PROJECT OVERVIEW', level=1)
        doc.add_paragraph(
            f'This Statement of Work ("SOW") defines the scope, deliverables, timeline, and terms '
            f'for {info["title"].lower()} services to be provided by [Contractor Name] ("Contractor") '
            f'for [Client Name] ("Client"). This SOW is governed by the terms of the Master Services '
            f'Agreement dated [Date] between the parties.'
        )
        
        doc.add_heading('1.1 Project Objectives', level=2)
        doc.add_paragraph('[Describe the primary business objectives this project aims to achieve. '
                         'What problem does this solve? What outcome does the client expect?]')
        
        doc.add_heading('1.2 Background', level=2)
        doc.add_paragraph('[Provide relevant context about the client\'s business, current situation, '
                         'and why this project is being undertaken.]')
        
        # Section 2: Scope of Work
        doc.add_heading('2. SCOPE OF WORK', level=1)
        doc.add_paragraph('The Contractor shall provide the following services:')
        for item in info["scope_items"]:
            doc.add_paragraph(item, style='List Bullet')
        
        doc.add_heading('2.1 Out of Scope', level=2)
        doc.add_paragraph('The following items are explicitly excluded from this SOW:')
        doc.add_paragraph('[Item 1 - e.g., ongoing maintenance beyond the support period]', style='List Bullet')
        doc.add_paragraph('[Item 2 - e.g., third-party licensing costs]', style='List Bullet')
        doc.add_paragraph('[Item 3 - e.g., content/copy creation by contractor]', style='List Bullet')
        
        doc.add_heading('2.2 Assumptions', level=2)
        doc.add_paragraph('[List key assumptions, e.g., client will provide content by X date, '
                         'client has existing hosting, etc.]', style='List Bullet')
        
        # Section 3: Deliverables
        doc.add_heading('3. DELIVERABLES', level=1)
        table = doc.add_table(rows=len(info["deliverables"])+1, cols=3)
        table.style = 'Light Shading Accent 1'
        headers = table.rows[0].cells
        headers[0].text = '#'
        headers[1].text = 'Deliverable'
        headers[2].text = 'Description'
        for h in headers:
            for p in h.paragraphs:
                for r in p.runs:
                    r.bold = True
        for i, (name, desc) in enumerate(info["deliverables"]):
            row = table.rows[i+1].cells
            row[0].text = str(i+1)
            row[1].text = name
            row[2].text = desc
        
        doc.add_paragraph()
        doc.add_heading('3.1 Acceptance Criteria', level=2)
        doc.add_paragraph('Each deliverable will be reviewed by the Client within [5] business days '
                         'of submission. Deliverables are considered accepted if no written objection '
                         'is received within the review period.')
        
        # Section 4: Timeline
        doc.add_heading('4. PROJECT TIMELINE & MILESTONES', level=1)
        table = doc.add_table(rows=len(info["milestones"])+1, cols=4)
        table.style = 'Light Shading Accent 1'
        headers = table.rows[0].cells
        for j, h in enumerate(['Milestone', 'Timeline', 'Payment %', 'Status']):
            headers[j].text = h
            for p in headers[j].paragraphs:
                for r in p.runs:
                    r.bold = True
        for i, (name, timeline, pct) in enumerate(info["milestones"]):
            row = table.rows[i+1].cells
            row[0].text = name
            row[1].text = timeline
            row[2].text = pct
            row[3].text = "☐ Not Started"
        
        # Section 5: Payment Terms
        doc.add_heading('5. COMPENSATION & PAYMENT TERMS', level=1)
        doc.add_heading('5.1 Project Fee', level=2)
        doc.add_paragraph(f'Typical rate range for {info["title"].lower()} services: {info["rate"]}')
        doc.add_paragraph()
        
        table = doc.add_table(rows=4, cols=2)
        table.style = 'Light Shading Accent 1'
        fee_items = [
            ("Total Project Fee:", "$[Amount]"),
            ("Payment Schedule:", "Per milestone completion (see Section 4)"),
            ("Late Payment Fee:", "1.5% per month on overdue balances"),
            ("Expense Reimbursement:", "Pre-approved expenses at cost + 0% markup"),
        ]
        for i, (label, value) in enumerate(fee_items):
            table.rows[i].cells[0].text = label
            table.rows[i].cells[1].text = value
        
        doc.add_heading('5.2 Payment Method', level=2)
        doc.add_paragraph('Payments shall be made via [bank transfer / PayPal / check] within '
                         '[15] days of invoice date.')
        
        doc.add_heading('5.3 Additional Work', level=2)
        doc.add_paragraph('Any work beyond the scope defined in this SOW will require a written '
                         'Change Order signed by both parties before work begins. Additional work '
                         'will be billed at $[X]/hour.')
        
        # Section 6: Terms
        doc.add_heading('6. TERMS & CONDITIONS', level=1)
        
        doc.add_heading('6.1 Intellectual Property', level=2)
        doc.add_paragraph('Upon receipt of full payment, all deliverables and intellectual property '
                         'created under this SOW shall be transferred to the Client. Contractor retains '
                         'the right to display work in portfolio unless otherwise agreed.')
        
        doc.add_heading('6.2 Confidentiality', level=2)
        doc.add_paragraph('Both parties agree to keep confidential any proprietary information '
                         'shared during the course of this project for a period of [2] years '
                         'following project completion.')
        
        doc.add_heading('6.3 Termination', level=2)
        doc.add_paragraph('Either party may terminate this SOW with [30] days written notice. '
                         'In the event of termination, Client shall pay for all work completed '
                         'and expenses incurred up to the termination date.')
        
        doc.add_heading('6.4 Limitation of Liability', level=2)
        doc.add_paragraph('Contractor\'s total liability under this SOW shall not exceed the total '
                         'project fee. Neither party shall be liable for indirect, incidental, or '
                         'consequential damages.')
        
        doc.add_heading('6.5 Force Majeure', level=2)
        doc.add_paragraph('Neither party shall be liable for delays caused by circumstances beyond '
                         'their reasonable control, including but not limited to natural disasters, '
                         'pandemics, or government actions.')
        
        doc.add_page_break()
        
        # Section 7: Signatures
        doc.add_heading('7. SIGNATURES', level=1)
        doc.add_paragraph('By signing below, both parties agree to the terms outlined in this '
                         'Statement of Work.')
        doc.add_paragraph()
        
        # Client signature block
        doc.add_paragraph('CLIENT', style='Heading 3')
        doc.add_paragraph()
        p = doc.add_paragraph('_' * 50)
        doc.add_paragraph('Signature')
        doc.add_paragraph()
        p = doc.add_paragraph('_' * 50)
        doc.add_paragraph('Printed Name & Title')
        doc.add_paragraph()
        p = doc.add_paragraph('_' * 50)
        doc.add_paragraph('Date')
        
        doc.add_paragraph()
        
        # Contractor signature block
        doc.add_paragraph('CONTRACTOR', style='Heading 3')
        doc.add_paragraph()
        p = doc.add_paragraph('_' * 50)
        doc.add_paragraph('Signature')
        doc.add_paragraph()
        p = doc.add_paragraph('_' * 50)
        doc.add_paragraph('Printed Name & Title')
        doc.add_paragraph()
        p = doc.add_paragraph('_' * 50)
        doc.add_paragraph('Date')
        
        filepath = os.path.join(outdir, f'sow-{slug}.docx')
        doc.save(filepath)
        print(f"  ✓ {filepath}")
    
    print("✅ SOW templates complete")

# ============================================================
# 2. RESTAURANT INVENTORY
# ============================================================
def create_restaurant_inventory():
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side, numbers
    from openpyxl.utils import get_column_letter
    from openpyxl.formatting.rule import CellIsRule, DataBarRule
    
    wb = Workbook()
    outdir = os.path.join(BASE, "restaurant-inventory", "files")
    os.makedirs(outdir, exist_ok=True)
    
    header_font = Font(name='Calibri', bold=True, color='FFFFFF', size=12)
    header_fill = PatternFill(start_color='1A3C6E', end_color='1A3C6E', fill_type='solid')
    subheader_fill = PatternFill(start_color='D6E4F0', end_color='D6E4F0', fill_type='solid')
    subheader_font = Font(name='Calibri', bold=True, size=11)
    thin_border = Border(
        left=Side(style='thin'), right=Side(style='thin'),
        top=Side(style='thin'), bottom=Side(style='thin')
    )
    
    def style_header(ws, row, cols):
        for c in range(1, cols+1):
            cell = ws.cell(row=row, column=c)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = Alignment(horizontal='center', wrap_text=True)
            cell.border = thin_border
    
    def style_range(ws, start_row, end_row, cols):
        for r in range(start_row, end_row+1):
            for c in range(1, cols+1):
                cell = ws.cell(row=r, column=c)
                cell.border = thin_border
                cell.alignment = Alignment(wrap_text=True)
    
    # --- Tab 1: Inventory Tracker ---
    ws = wb.active
    ws.title = "Inventory Tracker"
    headers = ['Item Name', 'Category', 'Unit', 'Par Level', 'On Hand', 'Reorder?', 
               'Unit Cost', 'Total Value', 'Supplier', 'Last Updated']
    for c, h in enumerate(headers, 1):
        ws.cell(row=1, column=c, value=h)
    style_header(ws, 1, len(headers))
    
    sample_data = [
        ("All-Purpose Flour", "Dry Goods", "50 lb bag", 5, 3, None, 18.50, None, "US Foods", "2026-02-15"),
        ("Olive Oil (Extra Virgin)", "Oils", "Gallon", 4, 6, None, 22.00, None, "Sysco", "2026-02-15"),
        ("Chicken Breast", "Protein", "Case (40 lb)", 3, 1, None, 89.00, None, "US Foods", "2026-02-14"),
        ("Heavy Cream", "Dairy", "Quart", 10, 4, None, 4.50, None, "Local Dairy Co", "2026-02-15"),
        ("Roma Tomatoes", "Produce", "Case (25 lb)", 4, 2, None, 32.00, None, "Fresh Direct", "2026-02-15"),
        ("Parmesan Cheese", "Dairy", "Wheel (5 lb)", 2, 3, None, 45.00, None, "Specialty Foods", "2026-02-13"),
        ("Garlic", "Produce", "Case (5 lb)", 3, 1, None, 15.00, None, "Fresh Direct", "2026-02-15"),
        ("Salmon Fillet", "Protein", "Case (10 lb)", 2, 2, None, 125.00, None, "Ocean Fresh", "2026-02-14"),
        ("Spaghetti Pasta", "Dry Goods", "Case (20 lb)", 4, 5, None, 24.00, None, "US Foods", "2026-02-12"),
        ("Unsalted Butter", "Dairy", "Case (36 ct)", 3, 2, None, 85.00, None, "Sysco", "2026-02-15"),
    ]
    for r, row_data in enumerate(sample_data, 2):
        for c, val in enumerate(row_data, 1):
            if c == 6:  # Reorder formula
                ws.cell(row=r, column=c, value=f'=IF(E{r}<D{r},"⚠️ REORDER","OK")')
            elif c == 8:  # Total Value formula
                ws.cell(row=r, column=c, value=f'=E{r}*G{r}')
                ws.cell(row=r, column=c).number_format = '$#,##0.00'
            else:
                ws.cell(row=r, column=c, value=val)
            if c == 7 and val is not None:
                ws.cell(row=r, column=c).number_format = '$#,##0.00'
    
    style_range(ws, 2, 11, len(headers))
    
    # Conditional formatting for reorder
    red_fill = PatternFill(start_color='FFC7CE', end_color='FFC7CE', fill_type='solid')
    green_fill = PatternFill(start_color='C6EFCE', end_color='C6EFCE', fill_type='solid')
    ws.conditional_formatting.add('F2:F100', CellIsRule(operator='equal', formula=['"⚠️ REORDER"'], fill=red_fill))
    ws.conditional_formatting.add('F2:F100', CellIsRule(operator='equal', formula=['"OK"'], fill=green_fill))
    
    # Summary row
    ws.cell(row=13, column=6, value="TOTAL INVENTORY VALUE:").font = Font(bold=True, size=12)
    ws.cell(row=13, column=8, value="=SUM(H2:H11)").font = Font(bold=True, size=12)
    ws.cell(row=13, column=8).number_format = '$#,##0.00'
    
    for c in range(1, len(headers)+1):
        ws.column_dimensions[get_column_letter(c)].width = 18
    
    # --- Tab 2: Food Cost Calculator ---
    ws2 = wb.create_sheet("Food Cost Calculator")
    headers2 = ['Menu Item', 'Selling Price', 'Ingredient 1', 'Cost 1', 'Ingredient 2', 'Cost 2',
                'Ingredient 3', 'Cost 3', 'Ingredient 4', 'Cost 4', 'Total Food Cost', 
                'Food Cost %', 'Gross Profit', 'Target Cost %', 'Status']
    for c, h in enumerate(headers2, 1):
        ws2.cell(row=1, column=c, value=h)
    style_header(ws2, 1, len(headers2))
    
    menu_items = [
        ("Grilled Salmon", 28.00, "Salmon", 6.50, "Vegetables", 1.20, "Sauce", 0.80, "Rice", 0.40),
        ("Chicken Parmesan", 22.00, "Chicken", 3.50, "Pasta", 0.60, "Sauce", 0.90, "Cheese", 1.50),
        ("Caesar Salad", 14.00, "Romaine", 1.00, "Dressing", 0.50, "Croutons", 0.30, "Parmesan", 0.80),
        ("Filet Mignon", 42.00, "Beef Filet", 14.00, "Potato", 0.80, "Vegetables", 1.00, "Butter", 0.40),
        ("Pasta Primavera", 18.00, "Pasta", 0.50, "Vegetables", 2.00, "Cream", 0.80, "Cheese", 1.00),
        ("Fish Tacos", 16.00, "Fish", 3.00, "Tortillas", 0.40, "Slaw", 0.60, "Sauce", 0.30),
    ]
    for r, item in enumerate(menu_items, 2):
        for c, val in enumerate(item, 1):
            ws2.cell(row=r, column=c, value=val)
            if c in (2, 4, 6, 8, 10):
                ws2.cell(row=r, column=c).number_format = '$#,##0.00'
        # Total Food Cost
        ws2.cell(row=r, column=11, value=f'=D{r}+F{r}+H{r}+J{r}')
        ws2.cell(row=r, column=11).number_format = '$#,##0.00'
        # Food Cost %
        ws2.cell(row=r, column=12, value=f'=K{r}/B{r}')
        ws2.cell(row=r, column=12).number_format = '0.0%'
        # Gross Profit
        ws2.cell(row=r, column=13, value=f'=B{r}-K{r}')
        ws2.cell(row=r, column=13).number_format = '$#,##0.00'
        # Target
        ws2.cell(row=r, column=14, value=0.30)
        ws2.cell(row=r, column=14).number_format = '0%'
        # Status
        ws2.cell(row=r, column=15, value=f'=IF(L{r}>N{r},"⚠️ OVER TARGET","✅ GOOD")')
    
    style_range(ws2, 2, 7, len(headers2))
    ws2.conditional_formatting.add('L2:L100', CellIsRule(operator='greaterThan', formula=['0.35'], fill=red_fill))
    ws2.conditional_formatting.add('L2:L100', CellIsRule(operator='lessThanOrEqual', formula=['0.30'], fill=green_fill))
    
    for c in range(1, len(headers2)+1):
        ws2.column_dimensions[get_column_letter(c)].width = 16
    
    # --- Tab 3: Order Log ---
    ws3 = wb.create_sheet("Order Log")
    headers3 = ['Order Date', 'Supplier', 'PO Number', 'Items Ordered', 'Qty', 'Unit Cost', 
                'Line Total', 'Received Date', 'Received Qty', 'Discrepancy', 'Status', 'Notes']
    for c, h in enumerate(headers3, 1):
        ws3.cell(row=1, column=c, value=h)
    style_header(ws3, 1, len(headers3))
    
    orders = [
        ("2026-02-10", "US Foods", "PO-2026-001", "All-Purpose Flour", 5, 18.50, None, "2026-02-12", 5, None, "Received", ""),
        ("2026-02-10", "US Foods", "PO-2026-001", "Chicken Breast", 3, 89.00, None, "2026-02-12", 3, None, "Received", ""),
        ("2026-02-11", "Sysco", "PO-2026-002", "Olive Oil", 4, 22.00, None, "2026-02-13", 4, None, "Received", ""),
        ("2026-02-13", "Fresh Direct", "PO-2026-003", "Roma Tomatoes", 4, 32.00, None, "2026-02-14", 3, None, "Short", "1 case damaged"),
        ("2026-02-14", "Ocean Fresh", "PO-2026-004", "Salmon Fillet", 3, 125.00, None, "", 0, None, "Pending", "Expected 2/16"),
    ]
    for r, row_data in enumerate(orders, 2):
        for c, val in enumerate(row_data, 1):
            if c == 7:
                ws3.cell(row=r, column=c, value=f'=E{r}*F{r}')
                ws3.cell(row=r, column=c).number_format = '$#,##0.00'
            elif c == 10:
                ws3.cell(row=r, column=c, value=f'=E{r}-I{r}')
            elif c == 6:
                ws3.cell(row=r, column=c, value=val)
                ws3.cell(row=r, column=c).number_format = '$#,##0.00'
            else:
                ws3.cell(row=r, column=c, value=val)
    style_range(ws3, 2, 6, len(headers3))
    for c in range(1, len(headers3)+1):
        ws3.column_dimensions[get_column_letter(c)].width = 16
    
    # --- Tab 4: Waste Log ---
    ws4 = wb.create_sheet("Waste Log")
    headers4 = ['Date', 'Item', 'Category', 'Qty Wasted', 'Unit', 'Unit Cost', 'Waste Cost',
                'Reason', 'Logged By', 'Prevention Notes']
    for c, h in enumerate(headers4, 1):
        ws4.cell(row=1, column=c, value=h)
    style_header(ws4, 1, len(headers4))
    
    waste = [
        ("2026-02-12", "Heavy Cream", "Dairy", 3, "Quart", 4.50, None, "Expired", "Chef Mike", "Order smaller qty"),
        ("2026-02-13", "Romaine Lettuce", "Produce", 2, "Head", 2.50, None, "Wilted", "Prep Cook", "Better rotation"),
        ("2026-02-14", "Salmon", "Protein", 1.5, "lb", 12.50, None, "Overcooked", "Line Cook", "Training needed"),
        ("2026-02-15", "Bread Rolls", "Bakery", 12, "Each", 0.75, None, "Stale", "FOH", "Reduce par level"),
    ]
    for r, row_data in enumerate(waste, 2):
        for c, val in enumerate(row_data, 1):
            if c == 7:
                ws4.cell(row=r, column=c, value=f'=D{r}*F{r}')
                ws4.cell(row=r, column=c).number_format = '$#,##0.00'
            elif c == 6:
                ws4.cell(row=r, column=c, value=val)
                ws4.cell(row=r, column=c).number_format = '$#,##0.00'
            else:
                ws4.cell(row=r, column=c, value=val)
    style_range(ws4, 2, 5, len(headers4))
    
    r_summary = 8
    ws4.cell(row=r_summary, column=5, value="TOTAL WASTE THIS PERIOD:").font = Font(bold=True, size=12)
    ws4.cell(row=r_summary, column=7, value="=SUM(G2:G6)").font = Font(bold=True, size=12)
    ws4.cell(row=r_summary, column=7).number_format = '$#,##0.00'
    
    for c in range(1, len(headers4)+1):
        ws4.column_dimensions[get_column_letter(c)].width = 16
    
    # --- Tab 5: Supplier List ---
    ws5 = wb.create_sheet("Supplier List")
    headers5 = ['Supplier Name', 'Contact Person', 'Phone', 'Email', 'Address',
                'Account #', 'Payment Terms', 'Min Order', 'Delivery Days', 'Categories', 'Rating', 'Notes']
    for c, h in enumerate(headers5, 1):
        ws5.cell(row=1, column=c, value=h)
    style_header(ws5, 1, len(headers5))
    
    suppliers = [
        ("US Foods", "John Smith", "(555) 123-4567", "john@usfoods.com", "123 Distribution Way", "USF-12345", "Net 30", "$500", "Tue, Thu", "Dry Goods, Protein", "★★★★★", "Primary supplier"),
        ("Sysco", "Jane Doe", "(555) 234-5678", "jane@sysco.com", "456 Supply Rd", "SYS-67890", "Net 30", "$300", "Mon, Wed, Fri", "Oils, Dairy, Equipment", "★★★★", "Good pricing on dairy"),
        ("Fresh Direct", "Bob Fresh", "(555) 345-6789", "bob@freshdirect.com", "789 Farm Ln", "FD-11111", "Net 15", "$200", "Daily", "Produce", "★★★★★", "Best produce quality"),
        ("Ocean Fresh", "Sara Waters", "(555) 456-7890", "sara@oceanfresh.com", "321 Harbor Dr", "OF-22222", "Net 15", "$400", "Tue, Fri", "Seafood", "★★★★", "Order by noon day before"),
        ("Local Dairy Co", "Tom Miller", "(555) 567-8901", "tom@localdairy.com", "654 Pasture Ave", "LDC-33333", "COD", "$100", "Mon, Wed", "Dairy, Eggs", "★★★★★", "Farm fresh, premium"),
    ]
    for r, row_data in enumerate(suppliers, 2):
        for c, val in enumerate(row_data, 1):
            ws5.cell(row=r, column=c, value=val)
    style_range(ws5, 2, 6, len(headers5))
    for c in range(1, len(headers5)+1):
        ws5.column_dimensions[get_column_letter(c)].width = 18
    
    filepath = os.path.join(outdir, 'restaurant-inventory-tracker.xlsx')
    wb.save(filepath)
    print(f"  ✓ {filepath}")
    print("✅ Restaurant inventory complete")

# ============================================================
# 3. COMPLIANCE CHECKLISTS
# ============================================================
def create_compliance_checklists():
    from fpdf import FPDF
    
    outdir = os.path.join(BASE, "compliance-checklists", "files")
    os.makedirs(outdir, exist_ok=True)
    
    checklists = {
        "llc-formation": {
            "title": "LLC Formation Checklist",
            "subtitle": "Complete Guide to Forming Your Limited Liability Company",
            "sections": [
                ("Pre-Formation Planning", [
                    "Research LLC requirements in your state",
                    "Choose a unique business name",
                    "Check name availability with Secretary of State",
                    "Decide on member-managed vs. manager-managed structure",
                    "Determine member ownership percentages",
                    "Select a registered agent",
                ]),
                ("Filing & Registration", [
                    "Prepare Articles of Organization",
                    "File Articles with Secretary of State",
                    "Pay state filing fee ($50-$500 depending on state)",
                    "Obtain Certificate of Formation",
                    "Apply for EIN (Employer Identification Number) from IRS",
                    "Register for state tax accounts",
                ]),
                ("Operating Agreement", [
                    "Draft Operating Agreement (even if single-member)",
                    "Define member roles and responsibilities",
                    "Establish voting rights and procedures",
                    "Set profit/loss distribution methods",
                    "Include buyout and dissolution provisions",
                    "Have all members sign and date",
                ]),
                ("Post-Formation", [
                    "Open a business bank account",
                    "Obtain required business licenses and permits",
                    "Set up accounting system",
                    "Purchase business insurance",
                    "File initial annual report (if required by state)",
                    "Create corporate records binder",
                ]),
            ]
        },
        "business-license": {
            "title": "Business License Checklist",
            "subtitle": "Comprehensive Licensing & Permit Requirements",
            "sections": [
                ("Federal Requirements", [
                    "Determine if federal license is required for your industry",
                    "Check FTC regulations for your business type",
                    "Obtain federal tax ID (EIN)",
                    "Register with relevant federal agencies (FDA, FCC, ATF, etc.)",
                    "Apply for import/export licenses if applicable",
                ]),
                ("State Requirements", [
                    "Research state-specific business license requirements",
                    "Apply for general state business license",
                    "Obtain professional/occupational licenses if required",
                    "Register for state sales tax permit",
                    "Apply for state employer registration",
                    "Check for industry-specific state permits",
                ]),
                ("Local Requirements", [
                    "Apply for city/county business license",
                    "Obtain zoning permits and verify zoning compliance",
                    "Apply for building permits if renovating",
                    "Get signage permits",
                    "Obtain health department permits (food businesses)",
                    "Apply for fire department permits",
                    "Check for home occupation permits (home-based businesses)",
                ]),
                ("Renewal & Maintenance", [
                    "Create a license renewal calendar",
                    "Set reminders 60 days before each expiration",
                    "Budget for annual renewal fees",
                    "Keep copies of all licenses on file",
                    "Display required licenses at place of business",
                    "Update licenses when business information changes",
                ]),
            ]
        },
        "tax-registration": {
            "title": "Tax Registration Checklist",
            "subtitle": "Federal, State & Local Tax Compliance Guide",
            "sections": [
                ("Federal Tax Registration", [
                    "Apply for EIN at IRS.gov (SS-4 form)",
                    "Determine business tax classification (sole prop, partnership, S-corp, C-corp)",
                    "Register for EFTPS (Electronic Federal Tax Payment System)",
                    "Understand quarterly estimated tax payment schedule",
                    "Set up payroll tax withholding (if employees)",
                    "Register for unemployment tax (FUTA)",
                ]),
                ("State Tax Registration", [
                    "Register with state Department of Revenue",
                    "Obtain state sales tax permit/resale certificate",
                    "Register for state income tax withholding",
                    "Register for state unemployment insurance (SUTA)",
                    "Check for franchise tax requirements",
                    "Register for use tax if applicable",
                ]),
                ("Local Tax Registration", [
                    "Register for city/county business tax",
                    "Check for local sales tax requirements",
                    "Register for property tax (business personal property)",
                    "Check for local payroll/head taxes",
                    "Research special district taxes",
                ]),
                ("Ongoing Compliance", [
                    "Set up tax calendar with all filing deadlines",
                    "Implement bookkeeping system (QuickBooks, Xero, etc.)",
                    "Maintain organized records for all deductions",
                    "Schedule quarterly tax review with accountant",
                    "File annual information returns (1099s, W-2s)",
                    "Keep tax records for minimum 7 years",
                ]),
            ]
        },
        "employment-law": {
            "title": "Employment Law Compliance Checklist",
            "subtitle": "Federal & State Labor Law Requirements for Employers",
            "sections": [
                ("Hiring Process", [
                    "Create compliant job descriptions (ADA-friendly)",
                    "Use consistent, non-discriminatory interview questions",
                    "Conduct proper background checks (FCRA compliance)",
                    "Verify employment eligibility (I-9 form within 3 days of hire)",
                    "Properly classify workers (employee vs. independent contractor)",
                    "Collect W-4 and state withholding forms",
                ]),
                ("Required Postings & Notices", [
                    "Display Federal Minimum Wage poster (FLSA)",
                    "Display OSHA workplace safety poster",
                    "Display EEOC 'EEO is the Law' poster",
                    "Display FMLA poster (50+ employees)",
                    "Display USERRA military service poster",
                    "Display state-required labor law posters",
                    "Provide new hire notices as required by state",
                ]),
                ("Wage & Hour Compliance", [
                    "Pay at least federal/state minimum wage (whichever is higher)",
                    "Track and pay overtime correctly (1.5x after 40 hrs/week)",
                    "Classify exempt vs. non-exempt employees properly",
                    "Maintain accurate time records for all non-exempt employees",
                    "Comply with meal and rest break requirements",
                    "Follow final paycheck laws upon termination",
                ]),
                ("Policies & Documentation", [
                    "Create employee handbook",
                    "Establish anti-discrimination and harassment policies",
                    "Implement workplace safety program",
                    "Create leave policies (FMLA, sick leave, PTO)",
                    "Establish termination and discipline procedures",
                    "Maintain confidential personnel files",
                    "Document all performance reviews and disciplinary actions",
                ]),
            ]
        },
        "osha-safety": {
            "title": "OSHA Workplace Safety Checklist",
            "subtitle": "Occupational Safety & Health Compliance Requirements",
            "sections": [
                ("General Safety Requirements", [
                    "Develop written safety and health program",
                    "Designate safety officer/coordinator",
                    "Conduct initial workplace hazard assessment",
                    "Create emergency action plan",
                    "Establish fire prevention plan",
                    "Post OSHA 'Job Safety and Health' poster",
                ]),
                ("Training & Documentation", [
                    "Provide safety orientation for all new employees",
                    "Conduct annual safety training refreshers",
                    "Train employees on hazard communication (HazCom/GHS)",
                    "Maintain Safety Data Sheets (SDS) for all chemicals",
                    "Document all training sessions with sign-in sheets",
                    "Train employees on proper PPE use and maintenance",
                ]),
                ("Workplace Conditions", [
                    "Maintain clear walking and working surfaces",
                    "Ensure proper lighting in all work areas",
                    "Maintain adequate ventilation",
                    "Keep emergency exits clear and properly marked",
                    "Maintain fire extinguishers (monthly inspections)",
                    "Ensure electrical systems are safe and up to code",
                    "Provide adequate first aid supplies",
                ]),
                ("Recordkeeping", [
                    "Maintain OSHA 300 Log of injuries and illnesses",
                    "Complete OSHA 301 forms for each recordable incident",
                    "Post OSHA 300A summary (Feb 1-Apr 30 annually)",
                    "Report fatalities within 8 hours",
                    "Report hospitalizations/amputations/eye loss within 24 hours",
                    "Keep records for minimum 5 years",
                ]),
            ]
        },
        "ada-compliance": {
            "title": "ADA Compliance Checklist",
            "subtitle": "Americans with Disabilities Act Requirements for Businesses",
            "sections": [
                ("Physical Accessibility", [
                    "Provide accessible parking spaces (van-accessible included)",
                    "Ensure accessible route from parking to entrance",
                    "Install automatic doors or accessible door hardware",
                    "Maintain 36-inch minimum clear width for aisles",
                    "Provide accessible restrooms",
                    "Ensure proper ramp slopes (1:12 maximum)",
                    "Install tactile/braille signage",
                ]),
                ("Employment (Title I)", [
                    "Review job descriptions for essential functions",
                    "Train hiring managers on ADA interview requirements",
                    "Establish interactive process for accommodation requests",
                    "Document all accommodation requests and responses",
                    "Ensure medical information is kept confidential",
                    "Review policies for potential disability discrimination",
                ]),
                ("Digital Accessibility", [
                    "Audit website for WCAG 2.1 AA compliance",
                    "Ensure all images have alt text",
                    "Provide captions for video content",
                    "Ensure keyboard navigation works throughout site",
                    "Test with screen readers (JAWS, NVDA, VoiceOver)",
                    "Ensure sufficient color contrast (4.5:1 ratio)",
                    "Make forms accessible with proper labels",
                ]),
                ("Policies & Training", [
                    "Create ADA compliance policy",
                    "Train staff on disability awareness and etiquette",
                    "Establish grievance procedure for ADA complaints",
                    "Designate ADA coordinator",
                    "Conduct annual accessibility audits",
                    "Maintain documentation of all compliance efforts",
                ]),
            ]
        },
        "data-privacy-gdpr": {
            "title": "Data Privacy & GDPR Checklist",
            "subtitle": "Data Protection Compliance for Businesses",
            "sections": [
                ("Data Mapping & Assessment", [
                    "Conduct data inventory - what personal data do you collect?",
                    "Map data flows (collection, storage, processing, sharing)",
                    "Identify legal basis for each data processing activity",
                    "Conduct Data Protection Impact Assessment (DPIA) if high-risk",
                    "Identify all third parties who receive personal data",
                    "Document data retention periods for each data type",
                ]),
                ("Policies & Notices", [
                    "Create/update Privacy Policy (clear, plain language)",
                    "Implement Cookie Policy and consent banner",
                    "Create internal Data Protection Policy",
                    "Establish Data Breach Response Plan",
                    "Create Data Subject Access Request (DSAR) procedure",
                    "Develop Employee Privacy Notice",
                ]),
                ("Technical Measures", [
                    "Implement data encryption (at rest and in transit)",
                    "Enable access controls and principle of least privilege",
                    "Set up regular data backups",
                    "Implement audit logging for data access",
                    "Use pseudonymization/anonymization where possible",
                    "Conduct regular security assessments/penetration testing",
                ]),
                ("Ongoing Compliance", [
                    "Train all employees on data protection (annual refresher)",
                    "Appoint Data Protection Officer (DPO) if required",
                    "Review and update data processing agreements with vendors",
                    "Conduct annual privacy impact assessments",
                    "Test breach response plan annually",
                    "Keep records of all processing activities (Article 30)",
                ]),
            ]
        },
        "insurance-requirements": {
            "title": "Business Insurance Requirements Checklist",
            "subtitle": "Essential Coverage for Protecting Your Business",
            "sections": [
                ("Required Insurance", [
                    "Workers' Compensation (required in most states with employees)",
                    "Unemployment Insurance",
                    "Disability Insurance (required in some states: CA, HI, NJ, NY, RI)",
                    "Commercial Auto Insurance (if business-owned vehicles)",
                    "Professional Liability/E&O (required for some licensed professions)",
                ]),
                ("Essential Coverage", [
                    "General Liability Insurance ($1M-$2M recommended)",
                    "Commercial Property Insurance",
                    "Business Interruption Insurance",
                    "Product Liability Insurance (if selling products)",
                    "Cyber Liability Insurance",
                    "Employment Practices Liability Insurance (EPLI)",
                ]),
                ("Industry-Specific", [
                    "Professional Liability / Errors & Omissions",
                    "Directors & Officers (D&O) Insurance",
                    "Liquor Liability (if serving alcohol)",
                    "Builders Risk Insurance (construction)",
                    "Marine/Cargo Insurance (shipping goods)",
                    "Equipment Breakdown Insurance",
                ]),
                ("Policy Management", [
                    "Review all policies annually with insurance broker",
                    "Update coverage when business changes occur",
                    "Maintain certificates of insurance on file",
                    "Provide certificates to landlords/clients as needed",
                    "Compare quotes from 3+ insurers at renewal",
                    "Bundle policies (BOP) for potential savings",
                    "Keep claims history organized",
                ]),
            ]
        },
        "intellectual-property": {
            "title": "Intellectual Property Checklist",
            "subtitle": "Protecting Your Business's Intellectual Assets",
            "sections": [
                ("Trademarks", [
                    "Conduct trademark search (USPTO TESS database)",
                    "Search state trademark databases",
                    "Check domain name availability",
                    "File federal trademark application (USPTO)",
                    "Consider international trademark registration (Madrid Protocol)",
                    "Monitor for trademark infringement",
                    "Maintain trademark (file Section 8 & 15 declarations)",
                ]),
                ("Copyrights", [
                    "Identify all copyrightable works (website, marketing, code)",
                    "Register key works with U.S. Copyright Office",
                    "Use proper copyright notices (© Year Name)",
                    "Implement content usage/licensing policies",
                    "Address employee vs. contractor ownership (work for hire)",
                    "Maintain records of creation dates and authorship",
                ]),
                ("Trade Secrets", [
                    "Identify proprietary information and trade secrets",
                    "Implement Non-Disclosure Agreements (NDAs)",
                    "Restrict access on a need-to-know basis",
                    "Mark confidential documents appropriately",
                    "Implement technical safeguards (encryption, access controls)",
                    "Include confidentiality clauses in employment agreements",
                ]),
                ("Patents & General IP", [
                    "Evaluate inventions for patent eligibility",
                    "File provisional patent applications for early protection",
                    "Conduct freedom-to-operate analysis",
                    "Create IP inventory/registry",
                    "Establish IP ownership clauses in all contracts",
                    "Schedule annual IP portfolio review",
                    "Budget for IP maintenance and enforcement",
                ]),
            ]
        },
        "health-department": {
            "title": "Health Department Compliance Checklist",
            "subtitle": "Food Safety & Health Code Requirements",
            "sections": [
                ("Permits & Licensing", [
                    "Obtain health department operating permit",
                    "Ensure all food handlers have valid food handler cards",
                    "Designate Certified Food Protection Manager on staff",
                    "Post health permit in visible location",
                    "Schedule and pass initial health inspection",
                    "Understand your jurisdiction's inspection scoring system",
                ]),
                ("Food Safety Practices", [
                    "Implement HACCP (Hazard Analysis Critical Control Points) plan",
                    "Maintain proper food storage temperatures (refrigerator <=41 degrees F, freezer <=0 degrees F)",
                    "Monitor cooking temperatures for all proteins",
                    "Follow proper cooling procedures (135 degrees F to 70 degrees F in 2 hrs, to 41 degrees F in 4 hrs)",
                    "Prevent cross-contamination (separate cutting boards, storage)",
                    "Implement proper handwashing procedures and stations",
                    "Follow FIFO (First In, First Out) inventory rotation",
                ]),
                ("Facility Requirements", [
                    "Maintain 3-compartment sink for dishwashing",
                    "Ensure proper ventilation and hood systems",
                    "Maintain grease trap (clean per schedule)",
                    "Ensure adequate lighting in prep and storage areas",
                    "Maintain pest control program (licensed provider)",
                    "Keep floors, walls, and ceilings clean and in good repair",
                    "Provide adequate employee restroom facilities",
                ]),
                ("Documentation", [
                    "Maintain daily temperature logs (coolers, freezers, hot holding)",
                    "Keep cleaning schedule and logs",
                    "Document employee training records",
                    "Maintain pest control reports",
                    "Keep equipment maintenance records",
                    "Store past health inspection reports",
                    "Maintain allergen information and ingredient lists",
                ]),
            ]
        },
        "fire-safety": {
            "title": "Fire Safety Compliance Checklist",
            "subtitle": "Fire Prevention & Emergency Preparedness",
            "sections": [
                ("Fire Prevention Systems", [
                    "Install and maintain fire alarm system",
                    "Install sprinkler system (if required by code)",
                    "Provide adequate fire extinguishers (one per 3,000 sq ft min)",
                    "Install emergency lighting and exit signs",
                    "Install kitchen hood suppression system (restaurants)",
                    "Ensure fire doors are functional and self-closing",
                ]),
                ("Inspections & Maintenance", [
                    "Schedule annual fire inspection with local fire marshal",
                    "Monthly fire extinguisher visual inspections",
                    "Annual professional fire extinguisher servicing",
                    "Quarterly sprinkler system inspections",
                    "Annual fire alarm testing",
                    "Semi-annual kitchen hood cleaning (restaurants)",
                    "Annual emergency lighting battery tests",
                ]),
                ("Emergency Preparedness", [
                    "Develop written fire evacuation plan",
                    "Post evacuation maps on each floor",
                    "Designate assembly points outside building",
                    "Assign fire wardens for each area/floor",
                    "Conduct fire drills (quarterly recommended)",
                    "Maintain emergency contact list",
                    "Train employees on fire extinguisher use (PASS method)",
                ]),
                ("Compliance Documentation", [
                    "Maintain fire inspection certificates",
                    "Keep records of all inspections and maintenance",
                    "Document fire drills (date, time, participants, issues)",
                    "Maintain fire watch logs when systems are impaired",
                    "Keep copies of fire safety permits",
                    "Store equipment certifications and warranties",
                ]),
            ]
        },
        "workers-comp": {
            "title": "Workers' Compensation Checklist",
            "subtitle": "Workers' Comp Insurance & Claims Management",
            "sections": [
                ("Policy Setup", [
                    "Determine if workers' comp is required in your state",
                    "Classify employees with correct NCCI codes",
                    "Obtain workers' comp policy from approved insurer",
                    "Post required workers' comp notices in workplace",
                    "Provide new hires with workers' comp information",
                    "Understand your state's reporting requirements",
                ]),
                ("Injury Prevention", [
                    "Implement workplace safety program",
                    "Conduct regular safety inspections",
                    "Provide required PPE and safety equipment",
                    "Train employees on safe work practices",
                    "Establish return-to-work/light duty program",
                    "Maintain safety committee (if required)",
                ]),
                ("Claims Management", [
                    "Create injury reporting procedures",
                    "Provide first aid immediately after injury",
                    "File First Report of Injury within required timeframe",
                    "Notify insurance carrier promptly",
                    "Investigate all workplace injuries",
                    "Maintain communication with injured worker",
                    "Coordinate with claims adjuster",
                ]),
                ("Recordkeeping & Audits", [
                    "Maintain OSHA 300 log of injuries",
                    "Keep detailed records of all claims",
                    "Prepare for annual premium audit",
                    "Track experience modification rate (EMR)",
                    "Review loss runs annually",
                    "Maintain payroll records by employee classification",
                ]),
            ]
        },
        "payroll-setup": {
            "title": "Payroll Setup Checklist",
            "subtitle": "Complete Payroll System Implementation Guide",
            "sections": [
                ("Pre-Payroll Setup", [
                    "Obtain EIN from IRS",
                    "Register with state tax/revenue agency",
                    "Register with state unemployment agency",
                    "Determine pay periods (weekly, bi-weekly, semi-monthly, monthly)",
                    "Establish pay rates and salary structures",
                    "Choose payroll processing method (in-house, service, software)",
                ]),
                ("Employee Documentation", [
                    "Collect W-4 (federal withholding) from each employee",
                    "Collect state withholding forms",
                    "Complete I-9 (employment eligibility) within 3 days of hire",
                    "Collect direct deposit authorization forms",
                    "Obtain signed offer letters with compensation details",
                    "Set up employee records in payroll system",
                ]),
                ("Payroll Processing", [
                    "Set up time tracking system",
                    "Calculate gross pay (regular + overtime)",
                    "Withhold federal income tax",
                    "Withhold state/local income taxes",
                    "Withhold FICA (Social Security 6.2% + Medicare 1.45%)",
                    "Process pre-tax deductions (401k, health insurance, HSA)",
                    "Process post-tax deductions (garnishments, Roth 401k)",
                    "Calculate employer taxes (FICA match, FUTA, SUTA)",
                ]),
                ("Tax Deposits & Filing", [
                    "Deposit federal taxes per deposit schedule (monthly or semi-weekly)",
                    "Deposit state taxes per state requirements",
                    "File Form 941 (quarterly federal tax return)",
                    "File state quarterly wage reports",
                    "File Form 940 (annual FUTA return)",
                    "Issue W-2s to employees by January 31",
                    "File W-2s/W-3 with SSA by January 31",
                    "Issue 1099s to contractors by January 31",
                ]),
            ]
        },
        "record-keeping": {
            "title": "Business Record Keeping Checklist",
            "subtitle": "Document Retention & Organization Guide",
            "sections": [
                ("Financial Records (Keep 7 years)", [
                    "Income statements / Profit & Loss statements",
                    "Balance sheets",
                    "General ledger and journal entries",
                    "Bank statements and reconciliations",
                    "Credit card statements",
                    "Accounts receivable and payable records",
                    "Expense receipts and documentation",
                ]),
                ("Tax Records (Keep 7 years)", [
                    "Filed tax returns (federal, state, local)",
                    "Supporting schedules and worksheets",
                    "1099s and W-2s issued",
                    "Sales tax returns and records",
                    "Property tax records",
                    "Tax payment confirmations",
                    "Depreciation schedules",
                ]),
                ("Employee Records", [
                    "Personnel files (keep 7 years after termination)",
                    "I-9 forms (keep 3 years after hire or 1 year after termination)",
                    "Payroll records (keep 4 years)",
                    "Benefits enrollment forms",
                    "Performance reviews and disciplinary actions",
                    "Workers' comp claims (keep 10 years)",
                    "Training records",
                ]),
                ("Corporate Records (Keep permanently)", [
                    "Articles of Incorporation/Organization",
                    "Operating Agreement / Bylaws",
                    "Meeting minutes and resolutions",
                    "Stock/membership certificates",
                    "Contracts and agreements",
                    "Intellectual property registrations",
                    "Insurance policies",
                    "Permits and licenses",
                ]),
            ]
        },
        "annual-compliance": {
            "title": "Annual Compliance Checklist",
            "subtitle": "Year-End & Annual Business Compliance Requirements",
            "sections": [
                ("Q1: January - March", [
                    "Issue W-2s to employees (by Jan 31)",
                    "Issue 1099s to contractors (by Jan 31)",
                    "File Form 940 (annual FUTA - by Jan 31)",
                    "File Form 941 (Q4 - by Jan 31)",
                    "File state annual reports / franchise tax (varies by state)",
                    "Renew business licenses and permits",
                    "Review and update insurance policies",
                    "Post OSHA 300A summary (Feb 1 - Apr 30)",
                ]),
                ("Q2: April - June", [
                    "File business tax returns (by Mar 15 for S-corps/partnerships or Apr 15)",
                    "File Form 941 (Q1 - by Apr 30)",
                    "Make Q1 estimated tax payment (by Apr 15)",
                    "Conduct mid-year budget review",
                    "Review employee handbook and update policies",
                    "Schedule mid-year performance reviews",
                    "Renew any Q2-expiring licenses",
                ]),
                ("Q3: July - September", [
                    "File Form 941 (Q2 - by Jul 31)",
                    "Make Q2 estimated tax payment (by Jun 15)",
                    "Conduct workplace safety audit",
                    "Review data privacy practices",
                    "Begin open enrollment planning for benefits",
                    "Make Q3 estimated tax payment (by Sep 15)",
                    "Review and update emergency plans",
                ]),
                ("Q4: October - December", [
                    "File Form 941 (Q3 - by Oct 31)",
                    "Conduct open enrollment for employee benefits",
                    "Review year-end tax planning strategies with CPA",
                    "Make Q4 estimated tax payment (by Jan 15 of next year)",
                    "Order W-2 and 1099 forms",
                    "Verify employee information for year-end reporting",
                    "Review contracts up for renewal",
                    "Conduct annual compliance training for all employees",
                    "Back up all business records",
                    "Set goals and budget for upcoming year",
                ]),
            ]
        },
    }
    
    for slug, data in checklists.items():
        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=20)
        
        # Header bar
        pdf.set_fill_color(26, 60, 110)
        pdf.rect(0, 0, 210, 45, 'F')
        
        pdf.set_text_color(255, 255, 255)
        pdf.set_font('Helvetica', 'B', 22)
        pdf.set_y(10)
        pdf.cell(0, 12, data["title"], align='C', new_x="LMARGIN", new_y="NEXT")
        
        pdf.set_font('Helvetica', '', 11)
        pdf.cell(0, 8, data["subtitle"], align='C', new_x="LMARGIN", new_y="NEXT")
        
        pdf.set_y(50)
        pdf.set_text_color(0, 0, 0)
        
        # Info box
        pdf.set_fill_color(230, 238, 248)
        pdf.set_draw_color(26, 60, 110)
        pdf.rect(15, 50, 180, 25, 'FD')
        pdf.set_xy(20, 52)
        pdf.set_font('Helvetica', '', 9)
        pdf.cell(85, 6, 'Business Name: _______________________________')
        pdf.cell(85, 6, 'Completed By: _______________________________', new_x="LMARGIN", new_y="NEXT")
        pdf.set_x(20)
        pdf.cell(85, 6, 'Date Started: _________________________________')
        pdf.cell(85, 6, 'Date Completed: ______________________________', new_x="LMARGIN", new_y="NEXT")
        
        pdf.set_y(80)
        
        for section_idx, (section_title, items) in enumerate(data["sections"]):
            # Check if we need a new page
            needed_height = 15 + len(items) * 8
            if pdf.get_y() + needed_height > 270:
                pdf.add_page()
            
            # Section header
            pdf.set_fill_color(44, 95, 138)
            pdf.set_text_color(255, 255, 255)
            pdf.set_font('Helvetica', 'B', 12)
            pdf.cell(0, 10, f'  {section_idx+1}. {section_title}', fill=True, new_x="LMARGIN", new_y="NEXT")
            pdf.ln(2)
            
            pdf.set_text_color(0, 0, 0)
            pdf.set_font('Helvetica', '', 10)
            
            for item in items:
                if pdf.get_y() > 270:
                    pdf.add_page()
                x = pdf.get_x()
                y = pdf.get_y()
                # Checkbox
                pdf.set_draw_color(100, 100, 100)
                pdf.rect(x + 5, y + 1, 4, 4)
                pdf.set_x(x + 14)
                pdf.multi_cell(170, 6, item, new_x="LMARGIN", new_y="NEXT")
                pdf.ln(1)
            
            pdf.ln(3)
        
        # Footer with notes
        if pdf.get_y() > 230:
            pdf.add_page()
        pdf.ln(5)
        pdf.set_fill_color(245, 245, 245)
        pdf.set_font('Helvetica', 'B', 11)
        pdf.cell(0, 8, '  NOTES:', fill=True, new_x="LMARGIN", new_y="NEXT")
        for _ in range(5):
            pdf.ln(8)
            pdf.cell(0, 0, '_' * 95)
        
        # Disclaimer
        pdf.ln(15)
        pdf.set_font('Helvetica', 'I', 8)
        pdf.set_text_color(100, 100, 100)
        pdf.multi_cell(0, 4, 
            'DISCLAIMER: This checklist is provided for informational purposes only and does not constitute legal, '
            'tax, or professional advice. Requirements vary by state, industry, and business type. Consult with '
            'qualified professionals (attorney, CPA, insurance broker) for advice specific to your situation. '
            'Laws and regulations change frequently - verify current requirements before relying on this checklist.')
        
        filepath = os.path.join(outdir, f'checklist-{slug}.pdf')
        pdf.output(filepath)
        print(f"  ✓ {filepath}")
    
    print("✅ Compliance checklists complete")

# ============================================================
# 4. RENTAL PROPERTY MANAGER
# ============================================================
def create_rental_property():
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter
    from openpyxl.formatting.rule import CellIsRule
    
    wb = Workbook()
    outdir = os.path.join(BASE, "rental-property-manager", "files")
    os.makedirs(outdir, exist_ok=True)
    
    hdr_font = Font(name='Calibri', bold=True, color='FFFFFF', size=12)
    hdr_fill = PatternFill(start_color='1A3C6E', end_color='1A3C6E', fill_type='solid')
    thin = Border(left=Side(style='thin'), right=Side(style='thin'), top=Side(style='thin'), bottom=Side(style='thin'))
    money_fmt = '$#,##0.00'
    pct_fmt = '0.0%'
    
    def hdr(ws, row, headers):
        for c, h in enumerate(headers, 1):
            cell = ws.cell(row=row, column=c, value=h)
            cell.font = hdr_font
            cell.fill = hdr_fill
            cell.alignment = Alignment(horizontal='center', wrap_text=True)
            cell.border = thin
    
    def borders(ws, r1, r2, cols):
        for r in range(r1, r2+1):
            for c in range(1, cols+1):
                ws.cell(row=r, column=c).border = thin
    
    def col_widths(ws, widths):
        for i, w in enumerate(widths, 1):
            ws.column_dimensions[get_column_letter(i)].width = w
    
    # --- Property Dashboard ---
    ws = wb.active
    ws.title = "Property Dashboard"
    headers = ['Property Name', 'Address', 'Type', 'Units', 'Occupied', 'Vacancy Rate',
               'Monthly Income', 'Monthly Expenses', 'Net Income', 'Cap Rate', 'Property Value']
    hdr(ws, 1, headers)
    
    props = [
        ("Maple Street Duplex", "123 Maple St, Denver, CO", "Duplex", 2, 2, None, 3200, 1800, None, None, 380000),
        ("Oak Avenue Apartments", "456 Oak Ave, Denver, CO", "Multi-Family", 8, 7, None, 12800, 7200, None, None, 1200000),
        ("Pine Road House", "789 Pine Rd, Boulder, CO", "Single Family", 1, 1, None, 2100, 1200, None, None, 425000),
        ("Cedar Court Townhome", "321 Cedar Ct, Lakewood, CO", "Townhome", 1, 0, None, 0, 650, None, None, 350000),
        ("Elm Street Fourplex", "654 Elm St, Aurora, CO", "Multi-Family", 4, 4, None, 6400, 3600, None, None, 680000),
    ]
    for r, p in enumerate(props, 2):
        for c, v in enumerate(p, 1):
            if c == 6:
                ws.cell(row=r, column=c, value=f'=1-(E{r}/D{r})')
                ws.cell(row=r, column=c).number_format = pct_fmt
            elif c == 9:
                ws.cell(row=r, column=c, value=f'=G{r}-H{r}')
                ws.cell(row=r, column=c).number_format = money_fmt
            elif c == 10:
                ws.cell(row=r, column=c, value=f'=(I{r}*12)/K{r}')
                ws.cell(row=r, column=c).number_format = pct_fmt
            else:
                ws.cell(row=r, column=c, value=v)
                if c in (7, 8, 11):
                    ws.cell(row=r, column=c).number_format = money_fmt
    borders(ws, 2, 6, len(headers))
    
    # Totals
    ws.cell(row=8, column=1, value="PORTFOLIO TOTALS").font = Font(bold=True, size=12)
    for c, formula in [(4, '=SUM(D2:D6)'), (5, '=SUM(E2:E6)'), (7, '=SUM(G2:G6)'), (8, '=SUM(H2:H6)'), (9, '=SUM(I2:I6)')]:
        ws.cell(row=8, column=c, value=formula).font = Font(bold=True, size=12)
        if c in (7, 8, 9):
            ws.cell(row=8, column=c).number_format = money_fmt
    ws.cell(row=8, column=6, value='=1-(E8/D8)')
    ws.cell(row=8, column=6).number_format = pct_fmt
    
    col_widths(ws, [22, 28, 14, 8, 10, 12, 16, 16, 14, 10, 16])
    
    # --- Tenant Tracker ---
    ws2 = wb.create_sheet("Tenant Tracker")
    h2 = ['Property', 'Unit', 'Tenant Name', 'Phone', 'Email', 'Lease Start', 'Lease End',
           'Monthly Rent', 'Security Deposit', 'Lease Status', 'Days Until Expiry', 'Notes']
    hdr(ws2, 1, h2)
    
    tenants = [
        ("Maple Street Duplex", "A", "Sarah Johnson", "(555) 111-2222", "sarah@email.com", "2025-06-01", "2026-05-31", 1600, 1600, None, None, "Excellent tenant"),
        ("Maple Street Duplex", "B", "Mike Chen", "(555) 222-3333", "mike@email.com", "2025-09-01", "2026-08-31", 1600, 1600, None, None, "Has 1 cat"),
        ("Oak Avenue Apartments", "101", "Lisa White", "(555) 333-4444", "lisa@email.com", "2025-03-01", "2026-02-28", 1600, 1600, None, None, "Renewing"),
        ("Oak Avenue Apartments", "102", "James Brown", "(555) 444-5555", "james@email.com", "2025-07-15", "2026-07-14", 1600, 1600, None, None, ""),
        ("Oak Avenue Apartments", "201", "Emma Davis", "(555) 555-6666", "emma@email.com", "2025-01-01", "2025-12-31", 1600, 1600, None, None, "Month-to-month"),
        ("Pine Road House", "-", "Robert Taylor", "(555) 666-7777", "robert@email.com", "2025-04-01", "2026-03-31", 2100, 2100, None, None, "Great tenant, maintains yard"),
        ("Elm Street Fourplex", "1", "Amy Wilson", "(555) 777-8888", "amy@email.com", "2025-08-01", "2026-07-31", 1600, 1600, None, None, ""),
    ]
    for r, t in enumerate(tenants, 2):
        for c, v in enumerate(t, 1):
            if c == 10:
                ws2.cell(row=r, column=c, value=f'=IF(G{r}="","",IF(G{r}<TODAY(),"EXPIRED",IF(G{r}-TODAY()<60,"EXPIRING SOON","ACTIVE")))')
            elif c == 11:
                ws2.cell(row=r, column=c, value=f'=IF(G{r}="","",G{r}-TODAY())')
            else:
                ws2.cell(row=r, column=c, value=v)
                if c in (8, 9):
                    ws2.cell(row=r, column=c).number_format = money_fmt
    borders(ws2, 2, 8, len(h2))
    col_widths(ws2, [22, 6, 16, 16, 22, 12, 12, 14, 16, 16, 16, 22])
    
    red_fill = PatternFill(start_color='FFC7CE', end_color='FFC7CE', fill_type='solid')
    yellow_fill = PatternFill(start_color='FFEB9C', end_color='FFEB9C', fill_type='solid')
    ws2.conditional_formatting.add('J2:J50', CellIsRule(operator='equal', formula=['"EXPIRED"'], fill=red_fill))
    ws2.conditional_formatting.add('J2:J50', CellIsRule(operator='equal', formula=['"EXPIRING SOON"'], fill=yellow_fill))
    
    # --- Rent Roll ---
    ws3 = wb.create_sheet("Rent Roll")
    h3 = ['Property', 'Unit', 'Tenant', 'Monthly Rent', 'Jan', 'Feb', 'Mar', 'Apr', 'May', 'Jun',
           'Jul', 'Aug', 'Sep', 'Oct', 'Nov', 'Dec', 'Total Collected', 'Balance Due']
    hdr(ws3, 1, h3)
    
    rent_data = [
        ("Maple St Duplex", "A", "Sarah Johnson", 1600, "✓", "✓", "", "", "", "", "", "", "", "", "", ""),
        ("Maple St Duplex", "B", "Mike Chen", 1600, "✓", "✓", "", "", "", "", "", "", "", "", "", ""),
        ("Oak Ave Apts", "101", "Lisa White", 1600, "✓", "LATE", "", "", "", "", "", "", "", "", "", ""),
        ("Oak Ave Apts", "102", "James Brown", 1600, "✓", "✓", "", "", "", "", "", "", "", "", "", ""),
        ("Pine Rd House", "-", "Robert Taylor", 2100, "✓", "✓", "", "", "", "", "", "", "", "", "", ""),
        ("Elm St 4plex", "1", "Amy Wilson", 1600, "✓", "✓", "", "", "", "", "", "", "", "", "", ""),
    ]
    for r, rd in enumerate(rent_data, 2):
        for c, v in enumerate(rd, 1):
            ws3.cell(row=r, column=c, value=v)
            if c == 4:
                ws3.cell(row=r, column=c).number_format = money_fmt
        # Total and balance formulas
        ws3.cell(row=r, column=17, value=f'=COUNTIF(E{r}:P{r},"✓")*D{r}')
        ws3.cell(row=r, column=17).number_format = money_fmt
        ws3.cell(row=r, column=18, value=f'=(D{r}*12)-Q{r}')
        ws3.cell(row=r, column=18).number_format = money_fmt
    
    borders(ws3, 2, 7, len(h3))
    col_widths(ws3, [16, 6, 16, 14] + [6]*12 + [16, 14])
    
    # --- Maintenance Log ---
    ws4 = wb.create_sheet("Maintenance Log")
    h4 = ['Date Reported', 'Property', 'Unit', 'Reported By', 'Issue Description', 'Priority',
           'Assigned To', 'Date Completed', 'Cost', 'Status', 'Notes']
    hdr(ws4, 1, h4)
    
    maint = [
        ("2026-01-15", "Maple St Duplex", "A", "Sarah Johnson", "Leaky kitchen faucet", "Medium", "Bob's Plumbing", "2026-01-17", 185, "Complete", "Replaced washer"),
        ("2026-01-22", "Oak Ave Apts", "201", "Emma Davis", "Heater not working", "High", "HVAC Pros", "2026-01-23", 350, "Complete", "Replaced ignitor"),
        ("2026-02-01", "Pine Rd House", "-", "Robert Taylor", "Garage door opener broken", "Low", "Overhead Door Co", "2026-02-08", 275, "Complete", "New motor installed"),
        ("2026-02-10", "Elm St 4plex", "3", "Tenant", "Bathroom fan noisy", "Low", "", "", 0, "Open", "Scheduled for 2/20"),
        ("2026-02-14", "Oak Ave Apts", "102", "James Brown", "Dishwasher leaking", "Medium", "Appliance Fix", "", 0, "In Progress", "Parts ordered"),
    ]
    for r, m in enumerate(maint, 2):
        for c, v in enumerate(m, 1):
            ws4.cell(row=r, column=c, value=v)
            if c == 9:
                ws4.cell(row=r, column=c).number_format = money_fmt
    borders(ws4, 2, 6, len(h4))
    col_widths(ws4, [14, 18, 6, 16, 28, 10, 18, 14, 10, 12, 24])
    
    red_f = PatternFill(start_color='FFC7CE', end_color='FFC7CE', fill_type='solid')
    ws4.conditional_formatting.add('F2:F50', CellIsRule(operator='equal', formula=['"High"'], fill=red_f))
    
    # --- Expense Tracker ---
    ws5 = wb.create_sheet("Expense Tracker")
    h5 = ['Date', 'Property', 'Category', 'Description', 'Vendor', 'Amount', 'Payment Method', 'Tax Deductible', 'Receipt', 'Notes']
    hdr(ws5, 1, h5)
    
    expenses = [
        ("2026-01-01", "Maple St Duplex", "Mortgage", "Monthly mortgage payment", "First National Bank", 1450, "Auto-pay", "Partial (interest)", "✓", ""),
        ("2026-01-01", "Oak Ave Apts", "Mortgage", "Monthly mortgage payment", "Wells Fargo", 5800, "Auto-pay", "Partial (interest)", "✓", ""),
        ("2026-01-05", "All Properties", "Insurance", "Quarterly premium", "State Farm", 2400, "Check", "Yes", "✓", "Covers all 5 properties"),
        ("2026-01-15", "Maple St Duplex", "Repairs", "Kitchen faucet repair", "Bob's Plumbing", 185, "Credit Card", "Yes", "✓", ""),
        ("2026-01-22", "Oak Ave Apts", "Repairs", "Heater repair", "HVAC Pros", 350, "Check", "Yes", "✓", "Unit 201"),
        ("2026-02-01", "All Properties", "Management", "Property management fee", "Self", 0, "-", "-", "-", "Self-managed"),
        ("2026-02-01", "Pine Rd House", "Mortgage", "Monthly mortgage payment", "Chase Bank", 1850, "Auto-pay", "Partial (interest)", "✓", ""),
    ]
    for r, e in enumerate(expenses, 2):
        for c, v in enumerate(e, 1):
            ws5.cell(row=r, column=c, value=v)
            if c == 6:
                ws5.cell(row=r, column=c).number_format = money_fmt
    borders(ws5, 2, 8, len(h5))
    
    r_tot = 10
    ws5.cell(row=r_tot, column=5, value="TOTAL EXPENSES:").font = Font(bold=True, size=12)
    ws5.cell(row=r_tot, column=6, value="=SUM(F2:F8)").font = Font(bold=True)
    ws5.cell(row=r_tot, column=6).number_format = money_fmt
    col_widths(ws5, [12, 20, 14, 26, 20, 14, 14, 16, 8, 24])
    
    # --- P&L Calculator ---
    ws6 = wb.create_sheet("P&L Calculator")
    h6 = ['Category', 'Jan', 'Feb', 'Mar', 'Apr', 'May', 'Jun', 'Jul', 'Aug', 'Sep', 'Oct', 'Nov', 'Dec', 'Annual Total']
    hdr(ws6, 1, h6)
    
    # Income section
    ws6.cell(row=2, column=1, value="INCOME").font = Font(bold=True, size=12, color='1A3C6E')
    income_rows = [
        ("Rental Income", [24700, 24700, 0,0,0,0,0,0,0,0,0,0]),
        ("Late Fees", [0, 75, 0,0,0,0,0,0,0,0,0,0]),
        ("Other Income", [0, 0, 0,0,0,0,0,0,0,0,0,0]),
    ]
    r = 3
    for label, vals in income_rows:
        ws6.cell(row=r, column=1, value=label)
        for c, v in enumerate(vals, 2):
            ws6.cell(row=r, column=c, value=v)
            ws6.cell(row=r, column=c).number_format = money_fmt
        ws6.cell(row=r, column=14, value=f'=SUM(B{r}:M{r})')
        ws6.cell(row=r, column=14).number_format = money_fmt
        r += 1
    
    ws6.cell(row=r, column=1, value="TOTAL INCOME").font = Font(bold=True)
    for c in range(2, 15):
        col_letter = get_column_letter(c)
        ws6.cell(row=r, column=c, value=f'=SUM({col_letter}3:{col_letter}5)')
        ws6.cell(row=r, column=c).number_format = money_fmt
        ws6.cell(row=r, column=c).font = Font(bold=True)
    r += 2
    
    ws6.cell(row=r, column=1, value="EXPENSES").font = Font(bold=True, size=12, color='1A3C6E')
    r += 1
    expense_cats = ["Mortgage/Loans", "Property Tax", "Insurance", "Repairs & Maintenance", "Utilities", "Management Fees", "Advertising", "Legal & Professional", "Miscellaneous"]
    exp_start = r
    for cat in expense_cats:
        ws6.cell(row=r, column=1, value=cat)
        for c in range(2, 14):
            ws6.cell(row=r, column=c, value=0)
            ws6.cell(row=r, column=c).number_format = money_fmt
        ws6.cell(row=r, column=14, value=f'=SUM(B{r}:M{r})')
        ws6.cell(row=r, column=14).number_format = money_fmt
        r += 1
    
    # Fill in some sample expense data
    sample_exp = {"Mortgage/Loans": 9100, "Property Tax": 2200, "Insurance": 800, "Repairs & Maintenance": 500}
    for i, cat in enumerate(expense_cats):
        if cat in sample_exp:
            for c in range(2, 4):  # Jan, Feb
                ws6.cell(row=exp_start+i, column=c, value=sample_exp[cat])
    
    ws6.cell(row=r, column=1, value="TOTAL EXPENSES").font = Font(bold=True)
    for c in range(2, 15):
        col_letter = get_column_letter(c)
        ws6.cell(row=r, column=c, value=f'=SUM({col_letter}{exp_start}:{col_letter}{r-1})')
        ws6.cell(row=r, column=c).number_format = money_fmt
        ws6.cell(row=r, column=c).font = Font(bold=True)
    exp_total_row = r
    income_total_row = 6
    r += 2
    
    ws6.cell(row=r, column=1, value="NET OPERATING INCOME").font = Font(bold=True, size=12, color='006600')
    for c in range(2, 15):
        col_letter = get_column_letter(c)
        ws6.cell(row=r, column=c, value=f'={col_letter}{income_total_row}-{col_letter}{exp_total_row}')
        ws6.cell(row=r, column=c).number_format = money_fmt
        ws6.cell(row=r, column=c).font = Font(bold=True, size=12)
    
    col_widths(ws6, [22] + [12]*13)
    borders(ws6, 2, r, 14)
    
    filepath = os.path.join(outdir, 'rental-property-manager.xlsx')
    wb.save(filepath)
    print(f"  ✓ {filepath}")
    print("✅ Rental property manager complete")

# ============================================================
# 5. SALON STARTER KIT
# ============================================================
def create_salon_kit():
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter
    
    outdir = os.path.join(BASE, "salon-starter-kit", "files")
    os.makedirs(outdir, exist_ok=True)
    
    hdr_font = Font(name='Calibri', bold=True, color='FFFFFF', size=11)
    hdr_fill = PatternFill(start_color='8B3A62', end_color='8B3A62', fill_type='solid')  # Salon-appropriate plum
    thin = Border(left=Side(style='thin'), right=Side(style='thin'), top=Side(style='thin'), bottom=Side(style='thin'))
    money = '$#,##0.00'
    
    # --- Client Intake Form ---
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    
    title = doc.add_heading('NEW CLIENT INTAKE FORM', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x8B, 0x3A, 0x62)
    
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run('[Your Salon Name]')
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(0x8B, 0x3A, 0x62)
    
    doc.add_paragraph()
    doc.add_heading('Personal Information', level=2)
    fields = [
        "Full Name: _______________________________________________    Date: ________________",
        "Phone: ________________________    Email: ___________________________________________",
        "Address: _________________________________________________________________________",
        "Date of Birth: ___________________    Referred By: ____________________________________",
        "Preferred Contact Method:   ☐ Phone   ☐ Text   ☐ Email",
    ]
    for f in fields:
        doc.add_paragraph(f)
    
    doc.add_heading('Hair History', level=2)
    hair_fields = [
        "Current hair color:   ☐ Natural   ☐ Color-treated   Color: ____________________________",
        "Last chemical service: __________________________    Date: ___________________________",
        "Products currently using: _________________________________________________________",
        "Previous salon: __________________________________________________________________",
        "What do you love about your hair? _________________________________________________",
        "What would you like to change? ____________________________________________________",
    ]
    for f in hair_fields:
        doc.add_paragraph(f)
    
    doc.add_heading('Health & Allergy Information', level=2)
    doc.add_paragraph('Please check any that apply:')
    health = [
        "☐ Allergies to hair products    Specify: ___________________________________________",
        "☐ Sensitive scalp    ☐ Scalp conditions (psoriasis, eczema, etc.)",
        "☐ Pregnant or nursing    ☐ Currently on medication that affects hair",
        "☐ Metal allergies    ☐ Latex allergy",
        "☐ Recent surgery or health conditions: ____________________________________________",
    ]
    for h in health:
        doc.add_paragraph(h)
    
    doc.add_heading('Service Preferences', level=2)
    prefs = [
        "Services interested in:   ☐ Cut   ☐ Color   ☐ Highlights   ☐ Balayage   ☐ Blowout",
        "   ☐ Keratin Treatment   ☐ Extensions   ☐ Perms   ☐ Nails   ☐ Waxing   ☐ Facial",
        "Preferred stylist: ___________________________    ☐ No preference",
        "Best appointment times:   ☐ Morning   ☐ Afternoon   ☐ Evening   ☐ Weekends",
    ]
    for p in prefs:
        doc.add_paragraph(p)
    
    doc.add_paragraph()
    doc.add_heading('Consent & Policies', level=2)
    doc.add_paragraph(
        'I understand that a consultation may be required before certain services. I agree to provide '
        'accurate health information and understand that withholding information may affect service results. '
        'I have read and agree to the salon\'s cancellation policy (24-hour notice required).'
    )
    doc.add_paragraph()
    doc.add_paragraph('Signature: ________________________________________    Date: ________________')
    
    doc.save(os.path.join(outdir, 'client-intake-form.docx'))
    print(f"  ✓ client-intake-form.docx")
    
    # --- Service Menu Template ---
    doc2 = Document()
    doc2.styles['Normal'].font.name = 'Calibri'
    doc2.styles['Normal'].font.size = Pt(11)
    
    title = doc2.add_heading('SERVICE MENU', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x8B, 0x3A, 0x62)
    
    sub = doc2.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run('[Your Salon Name]')
    r.font.size = Pt(16)
    r.font.color.rgb = RGBColor(0x8B, 0x3A, 0x62)
    
    services = {
        "HAIRCUTS & STYLING": [
            ("Women's Haircut & Style", "$55 - $85", "Includes consultation, shampoo, cut, and blowout"),
            ("Men's Haircut", "$30 - $45", "Includes consultation, shampoo, and cut"),
            ("Children's Cut (12 & under)", "$25 - $35", "Fun, kid-friendly experience"),
            ("Bang Trim", "$15", "Quick trim between appointments"),
            ("Blowout & Style", "$45 - $65", "Shampoo, blow dry, and style"),
            ("Special Occasion Updo", "$75 - $125", "Prom, wedding, formal events"),
            ("Bridal Hair (trial included)", "$150 - $250", "Consultation, trial, and day-of styling"),
        ],
        "COLOR SERVICES": [
            ("Single Process Color", "$85 - $120", "All-over color, root touch-up"),
            ("Root Touch-Up", "$65 - $85", "Regrowth color only"),
            ("Partial Highlights", "$95 - $140", "Face-framing or partial foils"),
            ("Full Highlights", "$140 - $200", "Full head foils"),
            ("Balayage / Ombre", "$175 - $275", "Hand-painted, natural-looking color"),
            ("Color Correction", "$200+/hour", "Consultation required - price varies"),
            ("Gloss / Toner", "$35 - $55", "Shine treatment, tone adjustment"),
        ],
        "TREATMENTS": [
            ("Deep Conditioning Treatment", "$25 - $45", "Intensive moisture treatment"),
            ("Keratin Smoothing Treatment", "$250 - $400", "Frizz reduction, lasts 3-6 months"),
            ("Scalp Treatment", "$35 - $55", "Exfoliating and nourishing"),
            ("Olaplex Treatment", "$35 - $50", "Bond-building repair treatment"),
        ],
        "NAIL SERVICES": [
            ("Classic Manicure", "$30 - $40", "Shape, cuticle care, polish"),
            ("Gel Manicure", "$45 - $55", "Long-lasting gel polish"),
            ("Classic Pedicure", "$45 - $55", "Soak, exfoliation, polish"),
            ("Spa Pedicure", "$65 - $80", "Deluxe with mask and extended massage"),
            ("Dip Powder Nails", "$50 - $65", "Durable, chip-resistant finish"),
            ("Nail Art (per nail)", "$5 - $15", "Custom designs available"),
        ],
        "WAXING": [
            ("Eyebrow Wax", "$18 - $25", "Shaping and cleanup"),
            ("Lip or Chin Wax", "$12 - $18", "Quick and precise"),
            ("Full Face Wax", "$45 - $60", "Brows, lip, chin, sides"),
            ("Bikini Wax", "$45 - $65", "Classic bikini line"),
            ("Full Leg Wax", "$65 - $85", "Ankle to upper thigh"),
        ],
    }
    
    for category, items in services.items():
        doc2.add_heading(category, level=1)
        table = doc2.add_table(rows=len(items)+1, cols=3)
        table.style = 'Light Shading Accent 1'
        headers = table.rows[0].cells
        headers[0].text = 'Service'
        headers[1].text = 'Price'
        headers[2].text = 'Description'
        for h in headers:
            for p in h.paragraphs:
                for run in p.runs:
                    run.bold = True
        for i, (name, price, desc) in enumerate(items):
            row = table.rows[i+1].cells
            row[0].text = name
            row[1].text = price
            row[2].text = desc
        doc2.add_paragraph()
    
    doc2.add_paragraph()
    p = doc2.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Prices may vary based on hair length, thickness, and condition.\n'
                   'Consultation recommended for color correction and major changes.\n'
                   'All prices include a complimentary consultation.')
    r.font.italic = True
    r.font.size = Pt(9)
    
    doc2.save(os.path.join(outdir, 'service-menu-template.docx'))
    print(f"  ✓ service-menu-template.docx")
    
    # --- Appointment Tracker ---
    wb = Workbook()
    ws = wb.active
    ws.title = "Weekly Schedule"
    
    headers = ['Time', 'Monday', 'Tuesday', 'Wednesday', 'Thursday', 'Friday', 'Saturday']
    for c, h in enumerate(headers, 1):
        cell = ws.cell(row=1, column=c, value=h)
        cell.font = hdr_font
        cell.fill = hdr_fill
        cell.alignment = Alignment(horizontal='center')
        cell.border = thin
    
    times = []
    for hour in range(9, 20):
        for minute in ['00', '30']:
            h12 = hour if hour <= 12 else hour - 12
            ampm = 'AM' if hour < 12 else 'PM'
            times.append(f'{h12}:{minute} {ampm}')
    
    for r, t in enumerate(times, 2):
        ws.cell(row=r, column=1, value=t).font = Font(bold=True)
        ws.cell(row=r, column=1).border = thin
        for c in range(2, 8):
            ws.cell(row=r, column=c).border = thin
    
    # Sample appointments
    samples = {
        (2, 2): "Sarah J. - Cut & Color\nStyleist: Maria",
        (4, 2): "Mike C. - Men's Cut\nStylist: Jen",
        (6, 3): "Lisa W. - Balayage\nStylist: Maria",
        (3, 4): "Emma D. - Blowout\nStylist: Jen",
        (8, 5): "Amy W. - Highlights\nStylist: Maria",
        (2, 6): "Walk-in - Men's Cut\nStylist: Available",
    }
    for (r_off, c), val in samples.items():
        ws.cell(row=r_off+1, column=c, value=val).alignment = Alignment(wrap_text=True)
    
    ws.column_dimensions['A'].width = 12
    for c in range(2, 8):
        ws.column_dimensions[get_column_letter(c)].width = 22
    ws.sheet_properties.pageSetUpPr = None
    
    wb.save(os.path.join(outdir, 'appointment-tracker.xlsx'))
    print(f"  ✓ appointment-tracker.xlsx")
    
    # --- Daily Revenue Tracker ---
    wb2 = Workbook()
    ws = wb2.active
    ws.title = "Daily Revenue"
    
    headers = ['Date', 'Client Name', 'Service', 'Stylist', 'Service Price', 'Product Sales',
               'Tip', 'Payment Method', 'Total', 'Notes']
    for c, h in enumerate(headers, 1):
        cell = ws.cell(row=1, column=c, value=h)
        cell.font = hdr_font
        cell.fill = hdr_fill
        cell.alignment = Alignment(horizontal='center', wrap_text=True)
        cell.border = thin
    
    samples = [
        ("2026-02-16", "Sarah Johnson", "Cut & Color", "Maria", 120, 28.50, 25, "Credit Card", None, ""),
        ("2026-02-16", "Mike Chen", "Men's Cut", "Jen", 35, 0, 10, "Cash", None, ""),
        ("2026-02-16", "Lisa White", "Balayage", "Maria", 225, 45, 40, "Credit Card", None, "New client"),
        ("2026-02-16", "Emma Davis", "Blowout", "Jen", 55, 15.50, 12, "Venmo", None, ""),
        ("2026-02-16", "Walk-in", "Men's Cut", "Available", 35, 0, 7, "Cash", None, ""),
    ]
    for r, s in enumerate(samples, 2):
        for c, v in enumerate(s, 1):
            if c == 9:
                ws.cell(row=r, column=c, value=f'=E{r}+F{r}+G{r}')
            else:
                ws.cell(row=r, column=c, value=v)
            if c in (5, 6, 7):
                ws.cell(row=r, column=c).number_format = money
            ws.cell(row=r, column=c).border = thin
        ws.cell(row=r, column=9).number_format = money
    
    # Daily totals
    r = 8
    ws.cell(row=r, column=4, value="DAILY TOTALS:").font = Font(bold=True, size=12)
    for c, label in [(5, "Services"), (6, "Products"), (7, "Tips"), (9, "Grand Total")]:
        ws.cell(row=r, column=c, value=f'=SUM({get_column_letter(c)}2:{get_column_letter(c)}6)')
        ws.cell(row=r, column=c).number_format = money
        ws.cell(row=r, column=c).font = Font(bold=True)
    
    for c in range(1, 11):
        ws.column_dimensions[get_column_letter(c)].width = 16
    
    wb2.save(os.path.join(outdir, 'daily-revenue-tracker.xlsx'))
    print(f"  ✓ daily-revenue-tracker.xlsx")
    
    # --- Inventory Checklist ---
    wb3 = Workbook()
    ws = wb3.active
    ws.title = "Product Inventory"
    
    headers = ['Product Name', 'Brand', 'Category', 'Size', 'Retail Price', 'Cost',
               'Margin %', 'In Stock', 'Min Stock', 'Reorder?', 'Supplier', 'Last Ordered']
    for c, h in enumerate(headers, 1):
        cell = ws.cell(row=1, column=c, value=h)
        cell.font = hdr_font
        cell.fill = hdr_fill
        cell.alignment = Alignment(horizontal='center', wrap_text=True)
        cell.border = thin
    
    products = [
        ("Moisture Shampoo", "Olaplex", "Shampoo", "8.5 oz", 30, 15, None, 8, 4, None, "Beauty Dist Co", "2026-01-15"),
        ("Repair Conditioner", "Olaplex", "Conditioner", "8.5 oz", 30, 15, None, 6, 4, None, "Beauty Dist Co", "2026-01-15"),
        ("Heat Protectant Spray", "Chi", "Styling", "6 oz", 18, 8, None, 12, 6, None, "SalonCentric", "2026-02-01"),
        ("Texturizing Spray", "Bumble & Bumble", "Styling", "4 oz", 32, 16, None, 3, 4, None, "Beauty Dist Co", "2026-01-20"),
        ("Color Depositing Mask", "Moroccanoil", "Treatment", "6.7 oz", 34, 17, None, 5, 3, None, "SalonCentric", "2026-02-10"),
        ("Smoothing Serum", "Living Proof", "Styling", "1.7 oz", 29, 14, None, 7, 4, None, "SalonCentric", "2026-01-28"),
        ("Dry Shampoo", "Batiste", "Shampoo", "6.7 oz", 10, 4.50, None, 15, 8, None, "Amazon Business", "2026-02-05"),
        ("Hair Oil", "Moroccanoil", "Treatment", "3.4 oz", 48, 24, None, 4, 3, None, "SalonCentric", "2026-02-10"),
    ]
    for r, p in enumerate(products, 2):
        for c, v in enumerate(p, 1):
            if c == 7:
                ws.cell(row=r, column=c, value=f'=(E{r}-F{r})/E{r}')
                ws.cell(row=r, column=c).number_format = '0.0%'
            elif c == 10:
                ws.cell(row=r, column=c, value=f'=IF(H{r}<I{r},"⚠️ REORDER","OK")')
            else:
                ws.cell(row=r, column=c, value=v)
                if c in (5, 6):
                    ws.cell(row=r, column=c).number_format = money
            ws.cell(row=r, column=c).border = thin
    
    from openpyxl.formatting.rule import CellIsRule
    red_fill = PatternFill(start_color='FFC7CE', end_color='FFC7CE', fill_type='solid')
    green_fill = PatternFill(start_color='C6EFCE', end_color='C6EFCE', fill_type='solid')
    ws.conditional_formatting.add('J2:J50', CellIsRule(operator='equal', formula=['"⚠️ REORDER"'], fill=red_fill))
    ws.conditional_formatting.add('J2:J50', CellIsRule(operator='equal', formula=['"OK"'], fill=green_fill))
    
    for c in range(1, 13):
        ws.column_dimensions[get_column_letter(c)].width = 16
    
    wb3.save(os.path.join(outdir, 'inventory-checklist.xlsx'))
    print(f"  ✓ inventory-checklist.xlsx")
    print("✅ Salon starter kit complete")

# ============================================================
# RUN ALL
# ============================================================
if __name__ == '__main__':
    print("\n🏗️  Creating Freelancer SOW Bundle...")
    create_sow_templates()
    
    print("\n🍽️  Creating Restaurant Inventory Tracker...")
    create_restaurant_inventory()
    
    print("\n📋 Creating Compliance Checklists...")
    create_compliance_checklists()
    
    print("\n🏠 Creating Rental Property Manager...")
    create_rental_property()
    
    print("\n💇 Creating Salon Starter Kit...")
    create_salon_kit()
    
    print("\n✅ ALL PRODUCTS CREATED!")
